from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

def save_to_word(courses, filename):
    doc = Document()
    doc.add_heading('Course List', 0)
    
    for i, (course_name, udemy_link) in enumerate(courses, 1):
        doc.add_paragraph(f"{i}. {course_name} - {udemy_link}")

    doc.save(filename)
    print(f"Saved to Word: {filename}")

def save_to_pdf(courses, filename):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    c.setFont("Helvetica", 12)
    c.drawString(100, height - 50, 'Course List')

    y = height - 70
    for i, (course_name, udemy_link) in enumerate(courses, 1):
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        
        c.drawString(100, y, f"{i}. {course_name} - {udemy_link}")
        y -= 20

    c.save()
    print(f"Saved to PDF: {filename}")

def process_courses_and_get_links():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True) # show chrome if False
        page = browser.new_page()
        page.goto("https://www.real.discount/")

        course_index = 0
        courses = []
        while True:
            course_elements = page.query_selector_all('h3.ml-3')

            if not course_elements:
                load_more_button = page.query_selector('input.btn.btn-primary.mb-5')
                if load_more_button:
                    load_more_button.click()
                    page.wait_for_timeout(5000)
                    print("Clicked 'Load More' to load more courses.")
                    continue
                else:
                    print("No 'Load More' button found. Ending process.")
                    break

            valid_courses = []
            for course in course_elements:
                skip_text = course.inner_text().strip()
                if skip_text == "The latest in learning. Online courses from $14.99.":
                    continue
                valid_courses.append(course)

            if course_index >= len(valid_courses):
                load_more_button = page.query_selector('input.btn.btn-primary.mb-5')
                if load_more_button:
                    load_more_button.click()
                    page.wait_for_timeout(5000)
                    print("Clicked 'Load More' to load more courses.")
                    continue
                else:
                    print("No 'Load More' button found to continue.")
                    break

            try:
                course = valid_courses[course_index]
                course_name = course.inner_text().strip()
                
                course.click()
                page.wait_for_timeout(5000)

                get_coupon_link = page.query_selector('a:has(span:text("Get Coupon"))')
                if get_coupon_link:
                    udemy_link = get_coupon_link.get_attribute('href')
                    if udemy_link and 'www.udemy.com' in udemy_link:
                        courses.append((course_name, udemy_link))
                        print(f"Course '{course_name}' - Udemy Link: {udemy_link}")
                    else:
                        print(f"Course '{course_name}' - No link containing 'www.udemy.com' found.")
                else:
                    print(f"Course '{course_name}' - 'Get Coupon' section not found.")

                page.go_back()
                page.wait_for_timeout(5000)

                course_index += 1

                if len(courses) in {20, 40, 60}: # Feel free to change the number as you like
                    if len(courses) == 20:
                        save_to_word(courses, "courses_20.docx")
                        save_to_pdf(courses, "courses_20.pdf")
                    elif len(courses) == 40:
                        save_to_word(courses, "courses_40.docx")
                        save_to_pdf(courses, "courses_40.pdf")
                    elif len(courses) == 60:
                        save_to_word(courses, "courses_60.docx")
                        save_to_pdf(courses, "courses_60.pdf")

            except Exception as e:
                print(f"Error processing course '{course_name}': {e}")
                page.go_back()
                page.wait_for_timeout(5000)
        
                course_elements = page.query_selector_all('h3.ml-3')

        if len(courses) > 0:
            save_to_word(courses, "courses_final.docx")
            save_to_pdf(courses, "courses_final.pdf")

        browser.close()

process_courses_and_get_links()

